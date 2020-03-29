import time
from docx import Document
import os
import re
class CreateDoc:

    def create(self,data=[]):
        if not os.path.exists('C:/LinkedinDocs'):
            os.mkdir('C:/LinkedinDocs')
        for user in data:
            document = Document()
            currentTime = time.time()
            currentTime = str(currentTime).split('.')[1]
            document.add_heading(user.get('personal_info',"Error 1").get('name',"Error 2"))
            document.add_heading(user.get('personal_info',"Error 1").get('headline',"Error 4"),level=2)
            phone = user.get('personal_info',"Error 1").get('phone','Phone error')
            email = user.get('personal_info',"Error 1").get('email','email Error')
            table = document.add_table(rows=1, cols=2)
            cell1 = table.cell(0,0)
            cell1.text = str(phone)
            cell2 = table.cell(0,1)
            cell2.text = str(email)
            document.add_paragraph(user.get('personal_info',"Error 1").get('summary',"Error 5"))
            document.add_heading("Experience",level=2)
            for job in user.get('experiences',"Error 3").get('jobs',"Error 4"):
                document.add_paragraph(job.get('title'),style='List Bullet')
                document.add_paragraph(job.get('date_range'))
                document.add_paragraph(str(job.get('description')).replace("\n","").strip())
            document.add_heading("Skills",level=2)
            for skill in user.get('skills'):
                document.add_paragraph(skill.get('name'),style='List Bullet')
            document.add_heading("Education",level=2)
            for edu in user.get('experiences',"Error 3").get('education',"Error 4"):
                document.add_paragraph(str(edu.get('field_of_study'))+'\t'+str(edu.get('date_range')),style='List Bullet')
                document.add_paragraph(edu.get('degree',''))
                document.add_paragraph(edu.get('name',"None"))
            name = str(user.get('personal_info',"Error 1").get('name',"Error 2"))
            result = re.sub(r'[^a-zA-Z]', " ", name)
            result = re.sub(' +', ' ', result)
            document.save('C:/LinkedinDocs/'+result.strip()+'.docx')
